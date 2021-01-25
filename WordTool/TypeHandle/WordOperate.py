import docx
import xlwt
import os
from docx.shared import Pt
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from WordTool.TypeHandle.BaseHandle import *
from WordTool.Common.GlobalVar import *
from win32com.client import Dispatch


class WordOperate(BaseHandle):
    __slots__ = ('__word_path', '__table_data', "doc_object", "activate_table_list", "excle_data")
    B_value = ""

    def __init__(self, path=None):
        super(BaseHandle, self).__init__()
        self.__word_path = path
        self.__table_data = {}
        self.doc_object = None
        self.activate_table_list = []
        self.excle_data = []

    def get_table_data(self):
        return self.__table_data

    def get_rules(self):
        pass

    def set_rules(self, rule):
        pass

    def set_table_data(self, uid, data):
        self.__table_data[uid] = data

    def set_word_path(self, path):

        self.__word_path = ""
        self.__table_data = {}
        self.doc_object = None
        self.activate_table_list = []
        self.excle_data = []

        self.__word_path = path

    def get_word_path(self):
        return self.__word_path

    def clear_activate_table_list(self):
        return self.activate_table_list.clear()

    def get_activate_table_list(self):
        return self.activate_table_list

    def load_word(self):
        '''
        加载word获取文本对象
        '''

        try:
            self.doc_object = docx.Document(self.__word_path)
            return True
        except:
            return False

    def get_all_table(self):
        '''
        :param
        :return:word中表的总数并加载table对象至对象字典
        '''
        table_count = 0
        if (os.path.exists(self.__word_path)):
            if self.load_word() is False:
                self.doc_object = docx.Document(self.__word_path)
            for table_index, table in enumerate(self.doc_object.tables):
                self.set_table_data(table_index, table)
                table_count = table_index + 1
        return table_count

    def get_tables(self, start=0, end=-1):
        '''
        :param:开始位置
        :param:结束位置
        :return:获取指定table对象列表，不展开里面值的情况下
        '''

        self.clear_table_data_list()

        if self.table_data is not None:

            for index in range(start, end):
                self.table_data_list.append(self.table_data.get(index))

        else:

            if self.get_all_table() >= 1:

                self.get_tables(start, end)
            else:
                pass

        return self.table_data_list

    def get_tables(self, row=0, column=0, target_value=""):
        '''
        :param:row 行
        :param:column列
        :return:获取指定table对象列表，展开里面值的情况下或的匹配对象
        '''

        self.clear_activate_table_list()
        if len(self.get_table_data()) != 0:
            for table in self.get_table_data().values():
                try:
                    if table.cell(row, column).text == target_value:
                        self.activate_table_list.append(table)
                except:
                    return []
        else:
            if self.get_all_table() != 0:
                self.get_tables(row, column, target_value)
            else:
                pass

        return self.activate_table_list

    def set_activate_table_range(self, start=1, end=-1):
        length = len(self.activate_table_list)
        if (start > end):
            return
        elif (start < 1 or end < 1):
            return
        elif (length < end):
            end = length
        start -= 1
        self.activate_table_list = self.activate_table_list[start: end]

    def doc2docx(self, path):
        '''
        将doc后缀文件转换docx并返回文件地址
        '''

        filename = os.path.splitext(path)
        if filename[1] == ".doc":
            w = Dispatch('Word.Application')
            w.Visible = 0
            w.DisplayAlerts = 0
            doc = w.Documents.Open(path)
            newpath = os.path.splitext(path)[0] + '.docx'
            doc.SaveAs(newpath, 16)
            w.Quit()
            return newpath
        else:
            return path

    def table_relational_mapping(self, requiredcells, is_standard=True, standard_line=3, dynamic_line=6):
        '''
        requiredcells=[[row1，clo1],[row2，clo2]]
        提取的表位置字段到excle_data中
        '''

        if is_standard:
            number = 0
            cell_temp_data = []
            for table in self.activate_table_list:
                row_idx = len(table.rows)
                column_idx = len(table.columns)
                for index, _temp in enumerate(requiredcells):
                    try:
                        if (row_idx <= _temp[0] or column_idx <= _temp[1]):
                            cell_str = ""
                        else:
                            cell_str = table.cell(*_temp).text

                        cell_temp_data.append([number, index, cell_str])
                    except ValueError:
                        continue
                number += 1
            self.update_excle_data(cell_temp_data)
        else:
            self._dynamic_table_line(requiredcells, standard_line, dynamic_line)
        # 保存文档word
        # self.doc_object.save(self.get_word_path())

    def _dynamic_table_line(self, requiredcells, standard_line, dynamic_line):
        '''
        is_standard=False时生成数据exlce存储列表值
        '''

        number = 0
        cell_temp_data = []
        max_value = self._get_required_max_row(requiredcells)

        for table in self.activate_table_list:
            n = int((len(table.rows) - standard_line) / dynamic_line)

            if standard_line < max_value:

                for _index in range(n):
                    for index, _temp in enumerate(requiredcells):

                        if _temp[0] > standard_line:
                            cell_str = table.cell(requiredcells[index][0] + dynamic_line * _index,
                                                  requiredcells[index][1]).text

                        else:
                            cell_str = table.cell(requiredcells[index][0], requiredcells[index][1]).text

                        cell_temp_data.append([number, index, cell_str])
                    number += 1
            else:
                for index in range(len(requiredcells)):
                    cell_str = table.cell(requiredcells[index][0], requiredcells[index][1]).text
                    cell_temp_data.append([number, index, cell_str])
                number += 1

        self.update_excle_data(cell_temp_data)

    def _get_required_max_row(self, requiredcells):
        '''
        获取你准备要提取的单元格中最大的row值
        '''

        max_value = 0
        for value in requiredcells:
            if value[0] > max_value:
                max_value = value[0]
        return max_value

    def _table_boundary_check(self, table, row=-1, clo=-1):
        '''
        检查你准备回去的单元格是否否和规范，是否有越界
        '''

        if len(table.rows) >= row and table._column_count >= clo:
            return True
        else:
            return False

    def save_excle(self, label=[], path="数据提取结果表.xls", sheetname="数据提取"):
        '''
        保存数据到excl文件中
        '''

        writebook = xlwt.Workbook()
        sheet = writebook.add_sheet(sheetname)
        # 如果设置label，就将label保存在数据的第一行
        if label:
            for data in self.excle_data:
                data[0] = data[0] + 1
            for index, data in enumerate(label):
                self.excle_data.append([0, index, data])
        for data in self.excle_data:
            sheet.write(*data)
        writebook.save(path)

    def update_excle_data(self, date_list):
        '''
        更excle存储列表中的数据
        '''

        self.excle_data.clear()
        self.excle_data = date_list
        return self.excle_data

    def ispath(self):
        '''
        验证路径是否合法，返回bool值
        '''

        if os.path(self.wordpath):
            return True
        return False

    def mergetable(self, requiredcells):
        '''
        合并文档中存在的全部该项数据到一个表格中
        '''
        number = 0
        cell_temp_data = []
        for table in self.activate_table_list:
            for index_x in range(len(table.rows)):
                if index_x == 0:
                    continue
                for index, _temp in enumerate(requiredcells):
                    try:
                        cell_str = table.cell(index_x, _temp[1]).text
                        cell_temp_data.append([number, index, cell_str])

                    except ValueError:
                        continue
                number += 1
        self.update_excle_data(cell_temp_data)

    @staticmethod
    def iterable_all_cells(obj_table):
        '''
        遍历表中全部单元格，
        '''
        for row in obj_table.rows:
            for cell in row.cells:
                yield cell

    @staticmethod
    def table_all_paragraphs_style(obj_table, ContentStyle):
        '''
        遍历表中全部单元格,设置字体样式，
        '''

        cells = WordOperate.iterable_all_cells(obj_table)
        for cell in cells:
            WordOperate.set_cell_border(cell,
                                        top={"sz": 24},
                                        bottom={"sz": 24},
                                        left={"sz": 24},
                                        right={"sz": 24}, )
            paragraphs = cell.paragraphs
            WordOperate.table_paragraphs_style(paragraphs, ContentStyle)

    @staticmethod
    def table_all_paragraphs_replace(obj_table, source, tag):
        '''
        遍历表中全部单元格,替换原有字符，
        '''
        cells = WordOperate.iterable_all_cells(obj_table)
        for cell in cells:
            paragraphs = cell.paragraphs
            WordOperate.table_paragraphs_replace(paragraphs, source, tag)

    @staticmethod
    def iterable_all_runs(paragraphs):
        '''
        遍历paragraphs中全部run，
        '''
        for paragraph in paragraphs:
            for run in paragraph.runs:
                yield run

    @staticmethod
    def table_paragraphs_replace(paragraphs, source, tag):
        '''
        段落值替换
        '''
        runs = WordOperate.iterable_all_runs(paragraphs)
        for run in runs:
            if source in run.text:
                run.text = run.text.replace(source, tag)

    @staticmethod
    def table_paragraphs_style(paragraphs, style):
        '''
        设置段落样式：包含字体，大小等
        '''
        runs = WordOperate.iterable_all_runs(paragraphs)
        for run in runs:
            run.text = WordOperate.remove_blank_space(run.text)
            if hasattr(ContentStyle, 'font'):
                run.font.name = style.font
                run.font.size = Pt(style.size)
                run.element.rPr.rFonts.set(qn('w:eastAsia'), style.font)
            else:
                continue

    @staticmethod
    def del_table_row(table_obj, row_index, contain=None):
        '''
        index 为list类型时支持批量删除指定行号
        删除word中table指定行，,并依赖第一个单元格值进行判断
        '''
        if type(row_index) is list:
            row_index.sort(reverse=True)
            try:
                if len(table_obj.rows) < row_index[0]:
                    return
                for index in row_index:
                    row = table_obj.rows[index]  # 获取目标行
                    row._element.getparent().remove(row._element)
            finally:
                return
        try:
            row = table_obj.rows[row_index]  # 获取目标行
        except IndexError:
            return
        if row.cells[0].text == contain and contain is not None:
            row._element.getparent().remove(row._element)
        elif contain is None:
            row._element.getparent().remove(row._element)

    @staticmethod
    def del_table_colunm(table_obj, index, contain=None):
        '''
        删除word中table指定列,并依赖第一个单元格值进行判断
        '''
        try:
            column = table_obj.columns[index]  # 获取目标列
        except IndexError:
            return
        if column.cells[0].text == contain and contain is not None:
            column._element.getparent().remove(column._element)
        elif contain is None:
            column._element.getparent().remove(column._element)

    @staticmethod
    def align_table(table_obj, postion, is_adaptivewidth=True):
        '''
        postion:调整表的对齐方式(LEFT,CENTER,RIGHT)
        is_adaptivewidth:默认依据窗体自动调整表格
        '''
        invert_op = getattr(WD_TABLE_ALIGNMENT, postion, None)
        if callable(invert_op):
            invert_op(WD_TABLE_ALIGNMENT.path.parent_op)
        table_obj.alignment = invert_op

        if is_adaptivewidth:
            table_obj.autofit = True

    @staticmethod
    def set_cell_border(cell, **kwargs):
        """
        Set cell`s border
        Usage:
        set_cell_border(
            cell,
            top={"sz": 12, "val": "single", "color": "#FF0000", "space": "0"},
            bottom={"sz": 12, "color": "#00FF00", "val": "single"},
            left={"sz": 24, "val": "dashed", "shadow": "true"},
            right={"sz": 12, "val": "dashed"},
        )
        """
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()

        # check for tag existnace, if none found, then create one
        tcBorders = tcPr.first_child_found_in("w:tcBorders")
        if tcBorders is None:
            tcBorders = OxmlElement('w:tcBorders')
            tcPr.append(tcBorders)

        # list over all available tags
        for edge in ('left', 'top', 'right', 'bottom', 'insideH', 'insideV'):
            edge_data = kwargs.get(edge)
            if edge_data:
                tag = 'w:{}'.format(edge)

                # check for tag existnace, if none found, then create one
                element = tcBorders.find(qn(tag))
                if element is None:
                    element = OxmlElement(tag)
                    tcBorders.append(element)

                # looks like order of attributes is important
                for key in ["sz", "val", "color", "space", "shadow"]:
                    if key in edge_data:
                        element.set(qn('w:{}'.format(key)), str(edge_data[key]))

    def save(self, path="NewWord.docx"):
        '''
        保存word，到指定路径
        '''
        self.doc_object.save(path)

    @staticmethod
    def remove_blank_space(_str, position=None):
        '''
        None:去除全部空格
        1:去除左边和右边空格
        2:去除右边空格
        3:去除左边空格
        '''
        if position == 0:
            return _str.strip()
        elif position == 1:
            return _str.rstrip()
        elif position == 2:
            return _str.lstrip()
        else:
            return _str.replace(" ", "")

    def branch_run(self, run_param):
        if run_param == 0:
            self.set_word_path(RULES_HANDLER_INSTANCE.get_s_path())
            self.get_tables(*RULES_HANDLER_INSTANCE.get_table_source_cell())
            self.table_relational_mapping(RULES_HANDLER_INSTANCE.get_table_required_cells())
            content = RULES_HANDLER_INSTANCE.get_table_content()
            content_style = self.to_content_style(RULES_HANDLER_INSTANCE.get_table_style())
            for table in self.get_activate_table_list():
                if content[0] and content[1]:
                    self.table_all_paragraphs_replace(table, content[0], content[1])
                if content[2]:
                    self.table_all_paragraphs_replace(table, " ", "")
                if (content[3] < content[4]) and (content[3] > 0):
                    self.del_table_row(table, [range(content[3], content[4])])
                elif (content[3] == content[4]) and (content[3] > 0):
                    self.del_table_row(table, content[3])
                self.table_all_paragraphs_style(table, content_style)
                # self.align_table(table, "CENTER")
            self.save(RULES_HANDLER_INSTANCE.get_s_path())

    def to_content_style(self, style):
        # style = RULES_HANDLER_INSTANCE.get_table_style()
        if style[0] == EnumFontSize.Size3_5:
            ContentStyle.size = 12
        elif style[0] == EnumFontSize.Size4_0:
            ContentStyle.size = 14
        elif style[0] == EnumFontSize.Size4_5:
            ContentStyle.size = 9
        elif style[0] == EnumFontSize.Size5_0:
            ContentStyle.size = 10.5
        elif style[0] == EnumFontSize.SizeDefault:
            ContentStyle.size = 12
        if style[1] == EnumFont.Black:
            ContentStyle.font = u'黑体'
        elif style[1] == EnumFont.Song:
            ContentStyle.font = u'宋体'
        elif style[1] == EnumFont.Default:
            ContentStyle.font = u'宋体'
        if style[2] == EnumFontAlignment.Alignment_center:
            ContentStyle.alignment = WD_PARAGRAPH_ALIGNMENT
        elif style[2] == EnumFontAlignment.AlignmentDefault:
            ContentStyle.alignment = WD_PARAGRAPH_ALIGNMENT
        if style[4] == EnumFontBold.Bold_off:
            ContentStyle.is_bold = False
        elif style[4] == EnumFontBold.Bold_On:
            ContentStyle.is_bold = True
        elif style[4] == EnumFontBold.BoldDefault:
            ContentStyle.is_bold = False
        return ContentStyle

    # def set_tables_font(self):
    #     '''
    #     设置表中样式
    #     '''
    #
    #     Dellines_index = [1, 2, 11, 12, 6, 14]
    #     ContentStyle.font = u'宋体'
    #     ContentStyle.size = 12
    #     self.get_tables(0, 0, u'测试需求名称')
    #     for table in self.get_activate_table_list():
    #         self.table_all_paragraphs_replace(table, "测试", "天下")
    #         self.del_table_row(table, Dellines_index)
    #         self.table_all_paragraphs_style(table, ContentStyle)
    #         self.align_table(table, "CENTER")
    #     self.save("D:\\ykq\\code\\ykq\\WordTool\\data\\test.docx")

    # def template_38(self, path):
    #     '''
    #     38所文档需求跟踪
    #     '''
    #
    #     self.set_word_path(self.doc2docx(path))
    #
    #     self.get_tables(0, 0, "测试需求名称")
    #     # self.get_tables(0, 0, "抽象用例标识")
    #     requiredcells = [[1, 2], [1, 5], [4, 4], [5, 4]]
    #
    #     # self.table_relational_mapping(requiredcells, False, 3, 6)
    #     self.table_relational_mapping(requiredcells)
    #     # requiredcells = [[0, 3], [0, 7], [1, 7]]
    #     #
    #     # self.table_relational_mapping(requiredcells)
    #
    #     self.save_excle()
    #
    # def template_381(self, path):
    #     '''
    #     38所文档需求跟踪
    #     '''
    #
    #     self.set_word_path(self.doc2docx(path))
    #
    #     WordOperate.B_value = "测试需求名称"
    #
    #     self.get_tables(0, 0, WordOperate.B_value)
    #
    #     # requiredcells = [[1, 2], [1, 5], [4, 4], [5, 4]]
    #     #
    #     # self.table_relational_mapping(requiredcells, False, 3, 6)
    #
    #     requiredcells = [[3, 4], [4, 4], [1, 2], [1, 5], [2, 3], [5, 4], [8, 4]]
    #
    #     self.table_relational_mapping(requiredcells)
    #
    #     self.save_excle()
    #
    # def template_14(self, path):
    #
    #     self.set_word_path(self.doc2docx(path))
    #
    #     # self.get_tables(0, 0, "类型")
    #     #
    #     # requiredcells = [[1, 1], [1, 2]]
    #     #
    #     # self.mergetable(requiredcells)
    #     self.get_tables(0, 0, "类型")
    #
    #     requiredcells = [[1, 1], [1, 2]]
    #
    #     self.mergetable(requiredcells)
    #
    #     self.save_excle()
    #
    # def temlate_38_std(self):
    #
    #     WordOperate.B_value = u"抽象用例标识"
    #
    #     self.get_tables(0, 0, WordOperate.B_value)
    #
    #     requiredcells = [[0, 2], [0, 7], [0, 2], [0, 7]]
    #
    #     self.table_relational_mapping(requiredcells)
    #
    #     self.save_excle()

    def run(self, run_param=None):

        if run_param is None:
            super().run(run_param)
        else:
            self.branch_run(run_param)

    def work_stop(self):
        pass

    @staticmethod
    def word_macro(index):
        pass


if __name__ == '__main__':
    test = WordOperate()
    # test.template_14("D:\\ykq\\code\\ykq\\WordTool\\data\\123.doc")
    # test.addtabeltitle("D:\\ykq\\code\\ykq\\WordTool\\data\\微服务\\微服务集成框架软件测试计划.docx")
    # test.getheadertitile("D:\\ykq\\code\\ykq\\WordTool\\data\\微服务\\微服务集成框架软件需求规格说明.docx")

    # test.template_38("D:\\ykq\\code\\ykq\\WordTool\\data\\系统监控工作站软件测试说明V1.00.docx")

    test.template_38("D:\\ykq\\code\\ykq\\WordTool\\data\\系统监控服务器软件测试计划V1.00.docx")
    # test.temlate_38_std()
