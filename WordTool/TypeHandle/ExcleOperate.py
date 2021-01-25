import docx
import xlwt
import zipfile
import os
import shutil

'''读取word中的文本'''


class ExcleOperate(object):
    __slots__ = ('D_TableData', '__SWordPath', "__word_path", "__table_data", "doc_object")

    def __init__(self, path=None):
        self.__word_path = path
        self.__table_data = None
        self.doc_object = None

    @property
    def table_data(self):

        return self.__table_data

    # @property.setter
    # def table_data(self, uid, data):
    #
    #     self.__table_data[uid] = data

    def set_word_path(self, path):

        self.__word_path = path

    def get_word_path(self):

        return self.__S_WordPath

    '''加载word获取文本对象'''

    def load_word(self, path):

        if self.doc_object is None:

            return docx.Document(path)

        else:

            return self.doc_object

    def get_all_table(self):
        '''
        :param
        :return:word中表的总数并加载table对象至对象字典
        '''
        doc = self.load_word(self.get_word_path())
        for table_index, table in enumerate(doc.tables):
            self.table_data[table_index] = table
        return table_index + 1

    def get_tables(self, start=0, end=-1):
        '''
        :param:开始位置
        :param:结束位置
        :return:获取指定table对象列表，不展开里面值的情况下
        '''

        table_data_list = []
        if self.table_data is not None:

            for index in range(start, end):
                table_data_list.append(self.table_data.get(index))

        else:

            if self.get_all_table() >= 1:

                self.get_tables(start, end)
            else:
                pass

        return table_data_list

    def get_tables(self, row=0, column=0, target_value=""):

        '''
        :param:row 行
        :param:column列
        :return:获取指定table对象列表，展开里面值的情况下或的匹配对象
        '''

        table_data_list = []
        if self.table_data is not None:
            for table in self.table_data.values():
                try:
                    if table.cell(row, column).text == target_value:
                        table_data_list.append(table)
                except:
                    return []
        else:
            if self.get_all_table() >= 1:
                self.get_tables(row, column, target_value)
            else:
                pass
        return table_data_list

    def Pathvalidation(self):

        if os.path(self.wordpath):
            return True
        return False

    '''读取word中的table'''

    def gettable(self, path=None):

        writebook = xlwt.Workbook()  # 打开一个excel
        sheet = writebook.add_sheet('test')
        sheet.write(0, 1, "1")
        if path == None:
            path = self.wordpath
        if path:
            doc = docx.Document(path)
            number = 0
            for tableindex, table in enumerate(doc.tables):  # 遍历所有表格
                print('----table------:' + str(tableindex + 1))
                # print("行:" + str(len(table.rows)))
                # print("列:" + str(len(table.columns)))
                if table.cell(0, 0).text == "测试需求名称":
                    number = number + 1
                    print("number:" + str(number))

                    sheet.write(number, 0, table.cell(1, 2).text)
                    sheet.write(number, 1, table.cell(1, 5).text)
                    sheet.write(number, 2, table.cell(4, 4).text)
                    sheet.write(number, 3, table.cell(5, 4).text)

                    print("需求名称:" + table.cell(1, 2).text)
                    print("需求标识:" + table.cell(1, 5).text)
                    print("测试项名称:" + table.cell(4, 4).text)
                    print("测试项标识:" + table.cell(5, 4).text)
                    if len(table.rows) >= 10:
                        x = int((len(table.rows) - 10) / 6)
                        for index in range(x):
                            number = number + 1
                            print("number:" + str(number))

                            print("需求名称:" + table.cell(1, 2).text)
                            print("需求标识:" + table.cell(1, 5).text)
                            print("测试项名称:" + table.cell((4 + 6 * (index + 1)), 4).text)
                            print("测试项标识:" + table.cell((5 + 6 * (index + 1)), 4).text)

                            sheet.write(number, 0, table.cell(1, 2).text)
                            sheet.write(number, 1, table.cell(1, 5).text)
                            sheet.write(number, 2, table.cell((4 + 6 * (index + 1)), 4).text)
                            sheet.write(number, 3, table.cell((5 + 6 * (index + 1)), 4).text)

                # for rowindex,row in enumerate(table.rows):  # 遍历表格的所有行
                #
                #
                # print("列:"+str(len(row.cells)))
                # for lineindex,cell in enumerate(row.cells):
                #     print(rowindex,lineindex,cell.text, '\t')

        else:
            print("没有word地址")

        writebook.save('data/2.xls')

    def getinfo(self, wordfile):
        f = zipfile.ZipFile(wordfile, 'r')
        for filename in f.namelist():
            f.extract(filename)
            print(filename)

    def getpic(self, path, zip_path, tmp_path, store_path):
        '''
        :param path:源文件
        :param zip_path:docx重命名为zip
        :param tmp_path:中转图片文件夹
        :param store_path:最后保存结果的文件夹（需要手动创建）
        :return:
        '''
        '''=============将docx文件重命名为zip文件===================='''
        os.rename(path, zip_path)
        # 进行解压
        f = zipfile.ZipFile(zip_path, 'r')
        # 将图片提取并保存
        for file in f.namelist():
            f.extract(file, tmp_path)
        # 释放该zip文件
        f.close()
        '''=============将docx文件从zip还原为docx===================='''
        os.rename(zip_path, path)
        # 得到缓存文件夹中图片列表
        pic = os.listdir(os.path.join(tmp_path, 'word/media'))
        '''=============将图片复制到最终的文件夹中===================='''
        for i in pic:
            # 根据word的路径生成图片的名称
            new_name = path.replace('\\', '_')
            new_name = new_name.replace(':', '') + '_' + i
            shutil.copy(os.path.join(tmp_path + '/word/media', i), os.path.join(store_path, new_name))
        '''=============删除缓冲文件夹中的文件，用以存储下一次的文件===================='''
        for i in os.listdir(tmp_path):
            # 如果是文件夹则删除
            if os.path.isdir(os.path.join(tmp_path, i)):
                shutil.rmtree(os.path.join(tmp_path, i))

    def gettxt(self):
        file = docx.Document("国产平台微服务架构服务治理和开发工具软件需求规格说明书v1.0 - 副本.docx")
        print("段落数:" + str(len(file.paragraphs)))  # 段落数为13，每个回车隔离一段

        # 输出每一段的内容
        # for para in file.paragraphs:
        #     print(para.text)

        # 输出段落编号及段落内容
        for i in range(len(file.paragraphs)):
            if len(file.paragraphs[i].text.replace(' ', '')) > 4:
                print("第" + str(i) + "段的内容是：" + file.paragraphs[i].text)

    @staticmethod
    def save_excle(excle_data, label=[], path="数据提取结果表.xls", sheetname="数据提取"):
        """保存数据到excl文件中"""

        writebook = xlwt.Workbook()
        sheet = writebook.add_sheet(sheetname)
        # 如果设置label，就将label保存在数据的第一行
        if label:
            for data in excle_data:
                data[0] = data[0] + 1

            for index, data in enumerate(label):
                excle_data.append([0, index, data])
        for data in excle_data:
            sheet.write(*data)
        writebook.save(path)


if __name__ == '__main__':
    # 源文件
    # path = r'E:\dogcat\提取图片\log.docx'
    # # docx重命名为zip
    # zip_path = r'E:\dogcat\提取图片\log.zip'
    # # 中转图片文件夹
    # tmp_path = r'E:\dogcat\提取图片\tmp'
    # # 最后保存结果的文件夹
    # store_path = r'E:\dogcat\提取图片\测试'
    # m = getpic(path, zip_path, tmp_path, store_path)

    # writebook = xlwt.Workbook()  # 打开一个excel
    # sheet = writebook.add_sheet('test')
    # sheet.write(0, 1, "1")
    # writebook.save('data/1.xlsx')

    file = docx.Document("D:\\ykq\\code\\ykq\\WordTool\\data\\1.docx")
    print("段落数:" + str(len(file.paragraphs)))  # 段落数为13，每个回车隔离一段

    # 输出每一段的内容
    # for para in file.paragraphs:
    #     print(para.text)
    writebook = xlwt.Workbook()  # 打开一个excel
    sheet = writebook.add_sheet('test')
    # sheet.write(0, 0, "1")
    x = 0
    # 输出段落编号及段落内容
    for i in file.paragraphs:
        style_name = i.style.name
        if style_name.startswith('Heading'):
            print(style_name, i.text, sep=':')

            if ("(MS_" in i.text.replace(' ', '')):
                tr = i.text.replace(' ', '')

                tr1, tr2 = tr.split('(')
                tr2.replace(' )', '')

                sheet.write(x, 0, tr1)
                sheet.write(x, 1, tr2)
                x = x + 1

        # print(file.paragraphs[i].text.replace(' ', ''))
        # if len(file.paragraphs[i].text.replace(' ', '')) > 4:
        #     print("第" + str(i) + "段的内容是：" + file.paragraphs[i].text)

    writebook.save('2.xls')
